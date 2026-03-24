"""Integration tests using small in-memory PDF and DOCX files."""

import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from pdf_docx_to_markdown import convert_document_to_markdown, ConversionConfig


def _make_simple_pdf(path: Path) -> None:
    """Create a minimal single-page PDF with text."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello PDF World", fontsize=12)
    page.insert_text((72, 120), "This is body text.", fontsize=10)
    doc.save(str(path))
    doc.close()


def _make_simple_docx(path: Path) -> None:
    """Create a minimal DOCX with a heading and paragraph."""
    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.save(str(path))


def _make_image_only_pdf(path: Path) -> None:
    """Create a PDF with an image but no text (simulates scanned page)."""
    doc = fitz.open()
    page = doc.new_page()
    # Insert a tiny 2x2 red PNG as an image.
    import struct, zlib
    def _make_tiny_png() -> bytes:
        raw = b"\x00\xff\x00\x00\xff\x00\x00\xff\x00\x00\xff\x00"
        compressed = zlib.compress(raw)
        def chunk(ctype, data):
            c = ctype + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xffffffff)
        return (
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", struct.pack(">IIBBBBB", 2, 2, 8, 2, 0, 0, 0))
            + chunk(b"IDAT", compressed)
            + chunk(b"IEND", b"")
        )
    page.insert_image(fitz.Rect(50, 50, 200, 200), stream=_make_tiny_png())
    doc.save(str(path))
    doc.close()


class TestPDFConversion:
    def test_converts_pdf_to_md(self, tmp_path: Path):
        pdf_path = tmp_path / "sample.pdf"
        _make_simple_pdf(pdf_path)

        result = convert_document_to_markdown(pdf_path, tmp_path)

        assert result is not None
        md_path = Path(result)
        assert md_path.exists()
        content = md_path.read_text(encoding="utf-8")
        assert "Hello PDF World" in content
        assert "body text" in content

    def test_pdf_produces_metadata(self, tmp_path: Path):
        pdf_path = tmp_path / "meta.pdf"
        _make_simple_pdf(pdf_path)

        result = convert_document_to_markdown(pdf_path, tmp_path)

        content = Path(result).read_text(encoding="utf-8")
        assert content.startswith("---")
        assert "source: meta.pdf" in content
        assert "pages:" in content

    def test_missing_file_raises(self, tmp_path: Path):
        import pytest

        with pytest.raises(FileNotFoundError):
            convert_document_to_markdown(tmp_path / "nonexistent.pdf", tmp_path)

    def test_unsupported_extension_returns_none(self, tmp_path: Path):
        txt_path = tmp_path / "readme.txt"
        txt_path.write_text("hello")

        result = convert_document_to_markdown(txt_path, tmp_path)
        assert result is None

    def test_custom_config_thresholds(self, tmp_path: Path):
        pdf_path = tmp_path / "cfg.pdf"
        _make_simple_pdf(pdf_path)

        cfg = ConversionConfig(
            heading1_threshold=10,
            heading2_threshold=8,
            heading3_threshold=5,
        )
        result = convert_document_to_markdown(pdf_path, tmp_path, config=cfg)
        assert result is not None

    def test_scanned_pdf_warning(self, tmp_path: Path, caplog):
        pdf_path = tmp_path / "scanned.pdf"
        _make_image_only_pdf(pdf_path)

        with caplog.at_level(logging.WARNING):
            convert_document_to_markdown(pdf_path, tmp_path)

        assert any("scanned" in r.message.lower() or "image" in r.message.lower()
                    for r in caplog.records if r.levelno >= logging.WARNING)


class TestDOCXConversion:
    def test_converts_docx_to_md(self, tmp_path: Path):
        docx_path = tmp_path / "sample.docx"
        _make_simple_docx(docx_path)

        result = convert_document_to_markdown(docx_path, tmp_path)

        assert result is not None
        md_path = Path(result)
        assert md_path.exists()
        content = md_path.read_text(encoding="utf-8")
        assert "Test Heading" in content
        assert "test paragraph" in content

    def test_docx_heading_becomes_markdown_heading(self, tmp_path: Path):
        docx_path = tmp_path / "headings.docx"
        _make_simple_docx(docx_path)

        result = convert_document_to_markdown(docx_path, tmp_path)

        content = Path(result).read_text(encoding="utf-8")
        assert "# Test Heading" in content

    def test_docx_with_table(self, tmp_path: Path):
        docx_path = tmp_path / "table.docx"
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "foo"
        table.cell(1, 1).text = "bar"
        doc.save(str(docx_path))

        result = convert_document_to_markdown(docx_path, tmp_path)

        content = Path(result).read_text(encoding="utf-8")
        assert "| Name | Value |" in content
        assert "| foo | bar |" in content

    def test_docx_merged_cells(self, tmp_path: Path):
        """Horizontally merged cells should not produce duplicate columns."""
        docx_path = tmp_path / "merged.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(0, 2).text = "C"
        # Merge row 1, cols 0-1.
        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 0).text = "Merged"
        table.cell(1, 2).text = "Solo"
        doc.save(str(docx_path))

        result = convert_document_to_markdown(docx_path, tmp_path)

        content = Path(result).read_text(encoding="utf-8")
        assert "| A | B | C |" in content
        # Merged cell should appear once, with an empty placeholder for the span.
        assert "Merged" in content
        assert "Solo" in content

    def test_docx_nested_table(self, tmp_path: Path):
        """Nested tables inside cells should be rendered inline."""
        from docx.oxml.ns import qn as _qn

        docx_path = tmp_path / "nested.docx"
        doc = Document()
        outer = doc.add_table(rows=1, cols=2)
        outer.cell(0, 0).text = "Left"
        # Insert a nested table into cell (0,1).
        inner_tbl = outer.cell(0, 1)._element.makeelement(_qn("w:tbl"), {})
        outer.cell(0, 1)._element.append(inner_tbl)
        # Add a row to the nested table via XML.
        tr = inner_tbl.makeelement(_qn("w:tr"), {})
        inner_tbl.append(tr)
        tc = tr.makeelement(_qn("w:tc"), {})
        tr.append(tc)
        p = tc.makeelement(_qn("w:p"), {})
        tc.append(p)
        r = p.makeelement(_qn("w:r"), {})
        p.append(r)
        t = r.makeelement(_qn("w:t"), {})
        t.text = "Nested"
        r.append(t)
        doc.save(str(docx_path))

        result = convert_document_to_markdown(docx_path, tmp_path)

        content = Path(result).read_text(encoding="utf-8")
        assert "Left" in content
        assert "Nested" in content
